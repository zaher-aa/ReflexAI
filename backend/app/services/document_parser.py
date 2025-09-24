import os
import mimetypes
import logging
from pathlib import Path
from typing import Optional, Tuple
from fastapi import HTTPException

# Document parsing libraries
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from striprtf.striprtf import rtf_to_text
    RTF_AVAILABLE = True
except ImportError:
    RTF_AVAILABLE = False

try:
    from odf import text, teletype
    from odf.opendocument import load
    ODF_AVAILABLE = True
except ImportError:
    ODF_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

logger = logging.getLogger(__name__)

class DocumentParser:
    """
    Universal document parser supporting multiple file formats commonly used by writers.
    """
    
    SUPPORTED_EXTENSIONS = {
        '.txt': 'text/plain',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.doc': 'application/msword',
        '.pdf': 'application/pdf',
        '.rtf': 'application/rtf',
        '.odt': 'application/vnd.oasis.opendocument.text',
        '.html': 'text/html',
        '.htm': 'text/html',
        '.md': 'text/markdown',
        '.markdown': 'text/markdown'
    }
    
    def __init__(self):
        """Initialize the document parser with available libraries."""
        self.available_parsers = {
            '.txt': True,
            '.docx': DOCX_AVAILABLE,
            '.doc': False,  # Legacy DOC files not supported
            '.pdf': PDF_AVAILABLE,
            '.rtf': RTF_AVAILABLE,
            '.odt': ODF_AVAILABLE,
            '.html': HTML_AVAILABLE,
            '.htm': HTML_AVAILABLE,
            '.md': MARKDOWN_AVAILABLE,
            '.markdown': MARKDOWN_AVAILABLE
        }
        
        # Log available parsers
        available = [ext for ext, available in self.available_parsers.items() if available]
        logger.info(f"Document parser initialized. Available formats: {', '.join(available)}")
    
    def detect_file_type(self, file_path: str, filename: str) -> Tuple[str, str]:
        """
        Detect file type from extension and optionally MIME type.
        Returns (extension, mime_type)
        """
        # Get extension from filename
        file_ext = Path(filename).suffix.lower()
        
        # Check if extension is supported
        if file_ext not in self.SUPPORTED_EXTENSIONS:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file format: {file_ext}. Supported formats: {', '.join(self.SUPPORTED_EXTENSIONS.keys())}"
            )
        
        # Check if parser is available for this format
        if not self.available_parsers.get(file_ext, False):
            raise HTTPException(
                status_code=500,
                detail=f"Parser not available for {file_ext} files. Please install required dependencies."
            )
        
        mime_type = self.SUPPORTED_EXTENSIONS[file_ext]
        return file_ext, mime_type
    
    async def parse_document(self, file_path: str, filename: str) -> str:
        """
        Parse document content from various file formats.
        Returns plain text content.
        """
        try:
            file_ext, mime_type = self.detect_file_type(file_path, filename)
            
            logger.info(f"Parsing document: {filename} ({file_ext})")
            
            # Route to appropriate parser based on file extension
            if file_ext == '.txt':
                return await self._parse_txt(file_path)
            elif file_ext == '.docx':
                return await self._parse_docx(file_path)
            elif file_ext == '.pdf':
                return await self._parse_pdf(file_path)
            elif file_ext == '.rtf':
                return await self._parse_rtf(file_path)
            elif file_ext == '.odt':
                return await self._parse_odt(file_path)
            elif file_ext in ['.html', '.htm']:
                return await self._parse_html(file_path)
            elif file_ext in ['.md', '.markdown']:
                return await self._parse_markdown(file_path)
            else:
                raise HTTPException(
                    status_code=400,
                    detail=f"No parser implementation for {file_ext}"
                )
                
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Error parsing document {filename}: {e}")
            raise HTTPException(
                status_code=500,
                detail=f"Failed to parse document: {str(e)}"
            )
    
    async def _parse_txt(self, file_path: str) -> str:
        """Parse plain text files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin1') as f:
                return f.read()
    
    async def _parse_docx(self, file_path: str) -> str:
        """Parse Microsoft Word .docx files."""
        if not DOCX_AVAILABLE:
            raise HTTPException(status_code=500, detail="python-docx library not available")
        
        doc = Document(file_path)
        text_content = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(' | '.join(row_text))
        
        return '\n\n'.join(text_content)
    
    async def _parse_pdf(self, file_path: str) -> str:
        """Parse PDF files."""
        if not PDF_AVAILABLE:
            raise HTTPException(status_code=500, detail="PyPDF2 library not available")
        
        text_content = []
        
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    logger.warning(f"Could not extract text from page {page_num + 1}: {e}")
        
        if not text_content:
            raise HTTPException(status_code=400, detail="No text could be extracted from PDF")
        
        return '\n\n'.join(text_content)
    
    async def _parse_rtf(self, file_path: str) -> str:
        """Parse Rich Text Format files."""
        if not RTF_AVAILABLE:
            raise HTTPException(status_code=500, detail="striprtf library not available")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            rtf_content = f.read()
        
        return rtf_to_text(rtf_content)
    
    async def _parse_odt(self, file_path: str) -> str:
        """Parse OpenDocument Text files."""
        if not ODF_AVAILABLE:
            raise HTTPException(status_code=500, detail="odfpy library not available")
        
        doc = load(file_path)
        text_content = []
        
        # Extract all text nodes
        for element in doc.getElementsByType(text.P):
            paragraph_text = teletype.extractText(element)
            if paragraph_text.strip():
                text_content.append(paragraph_text)
        
        return '\n\n'.join(text_content)
    
    async def _parse_html(self, file_path: str) -> str:
        """Parse HTML files."""
        if not HTML_AVAILABLE:
            raise HTTPException(status_code=500, detail="beautifulsoup4 library not available")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Get text content
        text = soup.get_text()
        
        # Clean up whitespace and normalize text for better readability analysis
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        clean_text = ' '.join(chunk for chunk in chunks if chunk)
        
        # Ensure proper sentence endings for readability calculation
        import re
        clean_text = re.sub(r'\s+', ' ', clean_text)  # Normalize whitespace
        clean_text = re.sub(r'([.!?])\s*([A-Z])', r'\1 \2', clean_text)  # Ensure space after sentence endings
        text = clean_text
        
        return text
    
    async def _parse_markdown(self, file_path: str) -> str:
        """Parse Markdown files."""
        if not MARKDOWN_AVAILABLE:
            raise HTTPException(status_code=500, detail="markdown library not available")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Convert markdown to HTML, then extract plain text
        html = markdown.markdown(md_content)
        
        if HTML_AVAILABLE:
            soup = BeautifulSoup(html, 'html.parser')
            return soup.get_text()
        else:
            # Fallback: return raw markdown
            return md_content
    
    def get_supported_formats(self) -> dict:
        """Get list of supported formats and their availability."""
        return {
            ext: {
                'mime_type': self.SUPPORTED_EXTENSIONS[ext],
                'available': self.available_parsers[ext],
                'description': self._get_format_description(ext)
            }
            for ext in self.SUPPORTED_EXTENSIONS.keys()
        }
    
    def _get_format_description(self, ext: str) -> str:
        """Get human-readable description of file format."""
        descriptions = {
            '.txt': 'Plain text files',
            '.docx': 'Microsoft Word documents (2007+)',
            '.doc': 'Legacy Microsoft Word documents',
            '.pdf': 'PDF documents',
            '.rtf': 'Rich Text Format',
            '.odt': 'OpenDocument Text',
            '.html': 'HTML documents',
            '.htm': 'HTML documents',
            '.md': 'Markdown files',
            '.markdown': 'Markdown files'
        }
        return descriptions.get(ext, f'{ext} files')